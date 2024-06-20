from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from datetime import datetime
from utils.file_manager import get_saved_images
from docx2pdf import convert



def generate_report():
    
    #Crea un nuevo documento word
    doc = Document()

    #Configura los margenes del documento
    configurar_margenes(doc)

    #Obtiene una imagen para agregar el logotipo de empresa
    logo_path = 'files/logo3.png'

    #Obtiene las imagenes temporales envidas en chat
    rutas_imagenes = get_saved_images()

    #Obtiene el titulo del reporte obtenido por chat
    textoAsignacion = 'Instalación de purga'

    #Obtiene la fecha actual
    current_dateTime = datetime.now()
    fechaHoy = f"{current_dateTime.day:02d}/{current_dateTime.month:02d}/{current_dateTime.year}"

    #Agrega header del documento
    agregar_encabezado(doc, logo_path)

    #Agrega titulo y fecha en una tabla en el formato de reporte
    agregar_tabla_informacion(doc, textoAsignacion, fechaHoy)

    #Agrega un espacio para formato de documento
    doc.add_paragraph()

    #Agrega las imagenes dependiendo de la cantidad de imagenes
    agregar_tabla_con_imagenes(doc, rutas_imagenes)


    #Variables para el nombrado del archivo, en esta parte se puede agregar tambien el titulo de reporte
    docx_filename = 'temp/ReporteConEncabezadoYMargenes.docx'
    pdf_filename = 'temp/ReporteConEncabezadoYMargenes.pdf'

    #Se guarda el archivo formato word e inmediatamente se convierte a pdf
    guardar_documento(doc, docx_filename)
    convertir_a_pdf(docx_filename, pdf_filename)

    return docx_filename

def configurar_margenes(doc):
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.25)
        section.bottom_margin = Inches(0.25)
        section.left_margin = Inches(0.25)
        section.right_margin = Inches(0.25)
        section.header_distance = Inches(0.018)

def agregar_encabezado(doc, logo_path):
    header = doc.sections[0].header
    header_table = header.add_table(rows=1, cols=2, width=Inches(12))
    header_table.style = 'Light Shading Accent 1'

    cell_logo = header_table.cell(0, 0)
    paragraph_logo = cell_logo.paragraphs[0]
    paragraph_logo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    run_logo = paragraph_logo.add_run()
    run_logo.add_picture(logo_path, width=Inches(2.9))
    cell_logo.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    cell_info = header_table.cell(0, 1)
    paragraph_info = cell_info.paragraphs[0]
    paragraph_info.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = paragraph_info.add_run('REPORTE FOTOGRÁFICO\n COMERCIALIZADORA DE PRODUCTOS Y SERVICIOS\n DE REFRIGERACIÓN RECAR SA DE CV')
    run.bold = True
    run.font.name = 'Segoe UI'
    run.font.size = Pt(10.2)
    run.font.color.rgb = RGBColor(13, 85, 137)
    cell_info.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def agregar_tabla_informacion(doc, textoAsignacion, fechaHoy):
    info_table = doc.add_table(rows=1, cols=2)
    info_table.style = 'Medium List 1 Accent 1'

    hdr_cells = info_table.rows[0].cells

    p0 = hdr_cells[0].paragraphs[0]
    run0 = p0.add_run('Asignación: ' + textoAsignacion)
    run0.font.name = 'Segoe UI'
    run0.font.size = Pt(10)
    run0.bold = False
    run0.font.color.rgb = RGBColor(13, 85, 137)

    p1 = hdr_cells[1].paragraphs[0]
    run1 = p1.add_run('Fecha: ' + fechaHoy)
    run1.font.name = 'Segoe UI'
    run1.font.size = Pt(10)
    run1.font.color.rgb = RGBColor(13, 85, 137)
    run1.bold = False

    hdr_cells[0].width = Inches(8)
    hdr_cells[1].width = Inches(1.6)

    p0.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    hdr_cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    hdr_cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP

def agregar_tabla_con_imagenes(doc, rutas_imagenes):
    num_imagenes = len(rutas_imagenes)
    
    if num_imagenes <= 2:
        rows = 2
        cols = 2
    elif num_imagenes <= 4:
        rows = 3
        cols = 2
    else:
        rows = ((num_imagenes - 1) // 3) + 2
        cols = 3

    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Light List Accent 1'

    if num_imagenes <= 4:
        cell_width = Inches(3)
        cell_height = Inches(4)
    else:
        cell_width = Inches(1.5)
        cell_height = Inches(2.6)

    index = 0
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            if index < num_imagenes:
                img_path = rutas_imagenes[index]
                paragraph = cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(img_path, width=cell_width, height=cell_height)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            index += 1


def guardar_documento(doc, filename):
    doc.save(filename)

def convertir_a_pdf(input_file, output_file):
    
    convert(input_file)
    convert(input_file, output_file)
    # convert("temp/")
